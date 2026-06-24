import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_pitch_document():
    doc = docx.Document()

    # Configure Font Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10.5)

    # ------------------ TITLE PAGE / HEADER ------------------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(36)
    title_run = title_p.add_run("MARINE FLOW - SEAFOOD FG STORE ERP")
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(46, 139, 87) # Sea Green primary color

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("Enterprise Pitch & Feature Specification Document")
    sub_run.italic = True
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph().paragraph_format.space_after = Pt(24)

    # ------------------ EXECUTIVE SUMMARY ------------------
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "Marine Flow is a state-of-the-art, high-fidelity Enterprise Resource Planning (ERP) application designed specifically "
        "for the seafood finished goods (FG) cold storage and export logistics industry. It addresses the operational complexities "
        "of seafood warehousing—including shelf-life decay, custom repacking, FIFO inventory allocation, strict chain-of-custody "
        "maker-checker approval workflows, and mobile container loading scan verifications.\n\n"
        "By digitizing warehouse logistics from processing floor to shipping container, Marine Flow eliminates human error, "
        "avoids product quality claims, blocks double-selling of reserved stock, and ensures 100% downstream traceability."
    )

    # ------------------ INDUSTRY PAIN POINTS & SOLUTIONS ------------------
    doc.add_heading("2. Core Industry Pain Points & Solutions", level=1)
    doc.add_paragraph(
        "In seafood logistics, standard generic ERP platforms fall short due to the unique characteristics of perishable inventory "
        "and client-specific branding requirements. Marine Flow solves these concrete challenges directly:"
    )

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    # Configure Headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Industry Pain Point'
    hdr_cells[1].text = 'Marine Flow Solution'
    hdr_cells[0].paragraphs[0].runs[0].font.bold = True
    hdr_cells[1].paragraphs[0].runs[0].font.bold = True

    pain_points = [
        (
            "Double-Selling & Stock Reservation: Sales and marketing teams committing overlapping stock to different customers, leading to order shortfalls and customer relationship damage.",
            "Strict FIFO Allocation: The system automatically locks the oldest available stock (matching grade, variety, and packing) to active PO line items, changing carton status to 'Reserved'. Once reserved, cartons are completely isolated from other allocations and general dispatches."
        ),
        (
            "Quality & Freshness Decay: Perishable seafood spoiling in cold storage due to poor rotation, causing massive waste and product claims.",
            "Dynamic Freshness & Aging Warnings: Renders dynamic alerts directly on the live dashboard. Cartons are flagged with orange warnings at 14+ days and red critical alerts at 30+ days, instructing warehouse operators to prioritize older stock."
        ),
        (
            "Complex Repacking & Genealogy: Generic inventory must be repacked into branded cartons, making it difficult to maintain traceability and link child products back to original lots.",
            "Repacking Genealogy Engine: Automatically consumes parent MCs (marking them 'Repacked') and spawns new child MCs that inherit PO allocation and section metadata. Child cartons preserve parent IDs, enabling 100% complete lineage audits."
        ),
        (
            "Loading Errors & Shipping Discrepancies: Warehouse crews loading incorrect grades, unallocated cartons, or duplicate boxes into export containers.",
            "Mobile Scan Verification: Responsive loader view with persistent auto-focus. Crews scan carton barcodes/short codes directly in real-time, verifying matching PO status, container weight, loading limits, and scan history before dispatch."
        ),
        (
            "Rented Store Operational Blindspots: Third-party rented cold rooms are managed by offline entities, preventing the employment of internal staff on site.",
            "Sender-Initiated Acceptance: Authorized source store managers can verify and accept transfers on behalf of destination rented stores, preserving system workflow continuity."
        ),
    ]

    for pain, solution in pain_points:
        row_cells = table.add_row().cells
        row_cells[0].text = pain
        row_cells[1].text = solution

    doc.add_paragraph().paragraph_format.space_before = Pt(12)

    # ------------------ FUNCTIONAL MODULES ------------------
    doc.add_heading("3. Functional Module Breakdown", level=1)

    # 3.1 Role-Based Access Control
    doc.add_heading("3.1 Role-Based Access Control (RBAC)", level=2)
    doc.add_paragraph(
        "Marine Flow protects sensitive data by enforcing a double-layered authorization schema: role-based permissions and physical store isolation. Users only see and manage stock related to their explicitly assigned cold stores."
    )
    for role, desc in [
        ("Admin", "Full control over user management, settings, master data definitions, and global system toggles."),
        ("General Manager (GM)", "Global operational oversight, master data editing, and cross-store movement approval capabilities. Restricted from admin user-management."),
        ("Marketing Manager", "Full ownership of PO creation, allocation overrides, and customer dispatch approvals. Blocked from physical stock movements."),
        ("Store Manager", "Dashboard monitoring, stock logs, and transfer approval/acceptance isolation strictly within assigned physical warehouses."),
        ("Operator", "Primary data entry role for recording production inwarding, initiating transfer requests, and scanning dispatches. All moves require manager checker approval.")
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{role}: ").bold = True
        p.add_run(desc)

    # 3.2 Dashboard
    doc.add_heading("3.2 Live Stock Analytics Dashboard", level=2)
    doc.add_paragraph(
        "A visually rich analytics hub providing real-time operational insights for authorized operators:"
    )
    for title, desc in [
        ("Live Stock Aggregation Grid", "Displays real-time quantities grouped by variety, grade, type, and packing, with expandable warehouse breakdown details."),
        ("Full Container Load (FCL) Conversion", "Automatically converts master carton inventory counts into export 40ft container equivalents based on variety capacity ratios."),
        ("Store Capacity Utilization Bar", "Tracks physical volume constraints by aggregating master carton weights and comparing them against physical storage limits (in tons)."),
        ("Export & Filter Panel", "Permits instant exports of styled, zebra-striped Excel reports containing real-time stock and PO requirements.")
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{title}: ").bold = True
        p.add_run(desc)

    # 3.3 Stock Movements
    doc.add_heading("3.3 Stock Movements (Maker-Checker)", level=2)
    doc.add_paragraph(
        "To prevent data entry errors, the application runs a secure dual-verification process for all cargo shifts:"
    )
    for m, desc in [
        ("Inwarding", "Registers processing-floor products. Automatically issues sequential MC numbers and maps barcodes."),
        ("Inter-Store Transfers", "Maker-Checker workflow. Operator initiates; initiating manager approves (cartons enter virtual 'In Transit' state); destination manager verifies physical cargo and accepts."),
        ("Dispatched Sale / Repack", "Exits stock from storage, verifying order matching or packing destination, logging full audit-trail logs.")
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{m}: ").bold = True
        p.add_run(desc)

    # 3.4 Custom Repacking
    doc.add_heading("3.4 Custom Repacking & Custom Barcodes", level=2)
    doc.add_paragraph(
        "Generic Master Cartons can be dispatched to repacking and returned as customer-branded products. "
        "Newly entered branded cartons inherit parent PO allocation metadata automatically, and custom buyer barcodes "
        "are mapped to children sequentially, ensuring zero gaps in traceability."
    )

    # 3.5 Container Loading
    doc.add_heading("3.5 Container Loading Scan Verification", level=2)
    doc.add_paragraph(
        "A responsive web interface optimized for rugged mobile scan devices in freezing conditions. Crew members scan carton codes "
        "which are instantly verified against the shipping purchase order database, blocking wrong grade allocations or duplicates instantly. "
        "An interactive circular progress ring shows real-time load progress."
    )

    # 3.6 Base32 Sequential Short Code
    doc.add_heading("3.6 Base32 Sequential Short Code Marking Guides", level=2)
    doc.add_paragraph(
        "For sites without active barcode scanners, Marine Flow generates confusable-safe 3-letter Base32 sequential short codes "
        "(e.g., 'A2D' to 'A4F'). Operators can toggle the UI to 'Manual MC Selection Mode' to check scrollable checkboxes by code, "
        "and print marking sheets with custom range instructions."
    )

    # ------------------ TECHNICAL STACK ------------------
    doc.add_heading("4. Technical Stack & Deployment Reliability", level=1)
    for tech, desc in [
        ("Frontend Architecture", "React-based Next.js App Router providing rich interactive states, clean layouts, and desktop-to-mobile responsiveness."),
        ("Database Engine", "SQLite database utilizing WAL (Write-Ahead Logging) mode and foreign key constraints for fast, transactional, single-writer multi-reader performance."),
        ("Containerization & Orchestration", "Docker and Docker Compose stack combining Next.js with localized persistence, deployable anywhere with zero dependencies."),
        ("Date & Security Standards", "JWT session tokens, robust HTTP middleware guards, and confusable-proof timezone-neutral date conversions.")
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{tech}: ").bold = True
        p.add_run(desc)

    doc.save("/home/ubuntu/FGStore/fg-store/Marine_Flow_Pitch_Document.docx")
    print("Pitch document created successfully!")

if __name__ == "__main__":
    create_pitch_document()
